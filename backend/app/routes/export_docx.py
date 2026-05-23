import io, json, base64, tempfile, os
from fastapi import APIRouter, Header, Request
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

router = APIRouter(prefix="/export-docx", tags=["export"])

FONT_CN = "Microsoft YaHei"

# ========== Helpers ==========

def _add_run(p, text, bold=False, size=None, color=None, font=None, italic=False):
    run = p.add_run(text)
    run.bold = bold
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if font: run.font.name = font
    if italic: run.italic = True
    return run

def _add_heading_line(doc, text, size=13, color=(0x33, 0x33, 0x33), border=True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    if border:
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "4",
            qn("w:space"): "1", qn("w:color"): "{:02X}{:02X}{:02X}".format(*color),
        })
        pBdr.append(bottom)
        pPr.append(pBdr)
    return p

def _add_avatar(doc, avatar):
    if avatar and avatar.startswith("data:image"):
        try:
            _, b64 = avatar.split(",", 1)
            img_data = base64.b64decode(b64)
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp.write(img_data)
                tmp_path = tmp.name
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(tmp_path, width=Inches(0.8))
            os.unlink(tmp_path)
        except Exception:
            pass

# ========== Template generators ==========

def _gen_deco(doc, data):
    """Art Deco: navy header, brass accents"""
    # Header
    if data.get("name"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, data["name"], bold=True, size=18, color=(0x1A, 0x27, 0x44))
    if data.get("jobStatus"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, data["jobStatus"], size=10, color=(0x66, 0x66, 0x66))

    contacts = [v for v in [data.get("email", ""), data.get("phone", ""), data.get("birthPlace", "")] if v]
    if contacts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, " | ".join(contacts), size=9, color=(0x88, 0x88, 0x88))

    _add_info_line(doc, data)
    doc.add_paragraph()
    _add_content(doc, data, heading_color=(0x1A, 0x27, 0x44), divider_color=(0xC8, 0xA9, 0x51))


def _gen_swiss(doc, data):
    """Swiss: clean white, tech blue, two-column via tab stops"""
    _add_avatar(doc, data.get("avatar", ""))
    if data.get("name"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        _add_run(p, data["name"], bold=True, size=20, color=(0x0F, 0x0F, 0x0F))
    if data.get("jobStatus"):
        p = doc.add_paragraph()
        _add_run(p, data["jobStatus"], size=9, color=(0x99, 0x99, 0x99))

    contacts = [data.get("email", ""), data.get("phone", "")]
    if any(contacts):
        p = doc.add_paragraph()
        _add_run(p, " · ".join(filter(None, contacts)), size=9, color=(0x99, 0x99, 0x99))

    _add_info_line(doc, data)
    doc.add_paragraph()
    _add_content(doc, data, heading_color=(0x00, 0x55, 0xFF), divider_color=(0x00, 0x55, 0xFF))


def _gen_luxury(doc, data):
    """Luxury: warm cream, gold double lines"""
    _add_avatar(doc, data.get("avatar", ""))
    # Gold lines at top
    for _ in range(2):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        pPr = p._element.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single", qn("w:sz"): "8" if _ == 0 else "4",
            qn("w:space"): "1", qn("w:color"): "B8944C",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    if data.get("name"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, data["name"], size=18, color=(0x2C, 0x1F, 0x14))
    if data.get("jobStatus"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, data["jobStatus"], size=10, color=(0x8A, 0x7A, 0x6A))

    contacts = [v for v in [data.get("email", ""), data.get("phone", ""), data.get("birthPlace", "")] if v]
    if contacts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, " | ".join(contacts), size=9, color=(0x8A, 0x7A, 0x6A))

    _add_info_line(doc, data)
    doc.add_paragraph()
    _add_content(doc, data, heading_color=(0xB8, 0x94, 0x4C), divider_color=(0xB8, 0x94, 0x4C))


def _gen_wabisabi(doc, data):
    """Wabi-Sabi: ash white, indigo accents"""
    _add_avatar(doc, data.get("avatar", ""))
    if data.get("name"):
        p = doc.add_paragraph()
        _add_run(p, data["name"], size=20, color=(0x2D, 0x40, 0x59))
    if data.get("jobStatus"):
        p = doc.add_paragraph()
        _add_run(p, data["jobStatus"], size=9, color=(0x88, 0x88, 0x88))

    contacts = [data.get("email", ""), data.get("phone", ""), data.get("birthPlace", "")]
    if any(contacts):
        p = doc.add_paragraph()
        _add_run(p, " · ".join(filter(None, contacts)), size=9, color=(0x88, 0x88, 0x88))

    _add_info_line(doc, data)
    doc.add_paragraph()
    _add_content(doc, data, heading_color=(0x2D, 0x40, 0x59), divider_color=(0x2D, 0x40, 0x59))


def _add_info_line(doc, data):
    parts = []
    if data.get("gender"): parts.append(f"性别：{data['gender']}")
    if data.get("birthDate"): parts.append(f"出生年月：{data['birthDate']}")
    if data.get("currentCity"): parts.append(f"现居：{data['currentCity']}")
    if data.get("targetCity"): parts.append(f"期望城市：{data['targetCity']}")
    if parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_run(p, " | ".join(parts), size=9, color=(0x88, 0x88, 0x88))


def _add_content(doc, data, heading_color, divider_color):
    # Personal strengths
    ps = data.get("personalStrengths", "")
    if ps:
        _add_heading_line(doc, "个人优势", color=heading_color)
        p = doc.add_paragraph(ps)
        p.paragraph_format.space_after = Pt(4)

    # Education
    educations = data.get("educations", [])
    if any(e.get("school") for e in educations):
        _add_heading_line(doc, "教育背景", color=heading_color)
        for e in educations:
            if not e.get("school"): continue
            line = e["school"]
            if e.get("major"): line += f" · {e['major']}"
            p = doc.add_paragraph()
            _add_run(p, line, bold=True, size=10.5)
            sub = []
            if e.get("education"): sub.append(e["education"])
            if e.get("startYear") and e.get("endYear"): sub.append(f"{e['startYear']}–{e['endYear']}")
            elif e.get("endYear"): sub.append(f"{e['endYear']}年毕业")
            if sub:
                p = doc.add_paragraph(" | ".join(sub))
                p.runs[0].font.size = Pt(9)
                p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            if e.get("campusExperience"):
                p = doc.add_paragraph(e["campusExperience"])
                p.runs[0].font.size = Pt(9)

    # Skills
    skills = data.get("skills", [])
    if skills:
        _add_heading_line(doc, "技能", color=heading_color)
        p = doc.add_paragraph(", ".join(skills))

    # Work
    works = [w for w in data.get("workExperience", []) if w.get("company") or w.get("title")]
    if works:
        _add_heading_line(doc, "工作经历", color=heading_color)
        for w in works:
            line = w.get("company", "")
            if w.get("title"): line += f" — {w['title']}"
            if w.get("startDate") or w.get("endDate"): line += f"  ({w.get('startDate', '')}–{w.get('endDate', '')})"
            p = doc.add_paragraph()
            _add_run(p, line, bold=True, size=10.5)
            if w.get("description"):
                p = doc.add_paragraph(w["description"])
                p.runs[0].font.size = Pt(9.5)

    # Projects
    projects = [p for p in data.get("projects", []) if p.get("name")]
    if projects:
        _add_heading_line(doc, "项目经历", color=heading_color)
        for pj in projects:
            line = pj["name"]
            if pj.get("role"): line += f" — {pj['role']}"
            p = doc.add_paragraph()
            _add_run(p, line, bold=True, size=10.5)
            if pj.get("description"):
                p = doc.add_paragraph(pj["description"])
                p.runs[0].font.size = Pt(9.5)

    # Honors
    honors = [h for h in data.get("honors", []) if h.get("name")]
    if honors:
        _add_heading_line(doc, "所获荣誉", color=heading_color)
        for h in honors:
            line = h["name"]
            if h.get("date"): line += f"（{h['date']}）"
            p = doc.add_paragraph(line, style="List Bullet")
            p.runs[0].font.size = Pt(10)


TEMPLATES = {"deco": _gen_deco, "swiss": _gen_swiss, "luxury": _gen_luxury, "wabisabi": _gen_wabisabi}


@router.post("")
async def export_docx(request: Request, authorization: str = Header(None)):
    body = await request.json()
    data = body.get("data", {})
    template = body.get("template", "deco")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = FONT_CN
    style.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Generate based on selected template
    generator = TEMPLATES.get(template, _gen_deco)
    generator(doc, data)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=resume.docx"},
    )
